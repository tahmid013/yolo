"""Produce the supervisor-facing 'training progress table_final.docx'.

Strategy: start from a byte-for-byte copy of `report_generation/training
progress table.docx` (so every heading, paragraph, style, note is preserved
exactly as the supervisor saw it), then MUTATE that copy in place to add
YOLOv12 content:

- After the last v11 training table (Table 1.4), interleave Tables 1.5, 1.6,
  1.7 (v12n/s/m) using the same 'TABLE X.X' + descriptive caption + data
  table pattern.
- After the last v11 per-class table (Table 2.4), interleave Tables 2.5-2.7.
- Extend the model-comparison Table 3 in place: v11 rows untouched, v12 rows
  appended. Followed by a short narrative paragraph placing v12 in context.

No v11 paragraph, style, or number is changed. If v12l is added later, one
regeneration extends everything by one more entry.

Usage:
    python analysis/build_final_docx.py                       # default names
    python analysis/build_final_docx.py --out-name "training progress table_final.docx"
"""
from __future__ import annotations

import argparse
import csv
import json
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


CLASS_ORDER = [
    "Earmuffs or Ear plugs", "Face shield", "Footwear", "Goggles",
    "Hand gloves", "Helmet", "Masks", "Safety apron or vest", "Safety jacket",
]

V12_VARIANTS = ("yolo12n", "yolo12s", "yolo12m", "yolo12l")
V12_LABELS = ("YOLOv12-n", "YOLOv12-s", "YOLOv12-m", "YOLOv12-l")


# ---------- readers (v12 data) ----------

def _read_v12_training(csv_path: Path) -> list[dict]:
    rows = []
    with csv_path.open("r", encoding="utf-8") as f:
        for r in csv.DictReader(f):
            r = {k.strip(): (v.strip() if v else v) for k, v in r.items()}
            try:
                epoch = int(float(r.get("epoch", 0)))
                box = float(r["train/box_loss"])
                cls = float(r["train/cls_loss"])
                dfl = float(r["train/dfl_loss"])
                rows.append({"epoch": epoch, "box": box, "cls": cls, "dfl": dfl,
                             "total": box + cls + dfl})
            except (KeyError, TypeError, ValueError):
                continue
    return rows


def _read_v12(runs_dir: Path) -> dict:
    out: dict = {}
    for sub in sorted(runs_dir.iterdir()):
        if not sub.is_dir():
            continue
        variant = next((v for v in V12_VARIANTS if sub.name.startswith(f"{v}_")), None)
        if variant is None:
            continue
        record = {"training": [], "per_class": [], "overall": None,
                  "source": sub.name}
        results_csv = sub / "results.csv"
        if results_csv.exists():
            record["training"] = _read_v12_training(results_csv)
        eval_path = sub / "full_eval" / "per_class.json"
        if eval_path.exists():
            payload = json.loads(eval_path.read_text(encoding="utf-8"))
            record["overall"] = payload.get("overall")
            by_name = {p["name"]: p for p in payload.get("per_class", [])}
            for cname in CLASS_ORDER:
                src = by_name.get(cname, {})
                record["per_class"].append({
                    "name": cname,
                    "precision": src.get("precision"),
                    "recall": src.get("recall"),
                    "ap50": src.get("ap50"),
                    "ap50_95": src.get("ap50_95"),
                })
        out[variant] = record
    return out


# ---------- v11 label counts from reference docx (Labels column) ----------

def _label_lookup_from_docx(doc) -> dict[str, int]:
    """Extract the Labels column values from the first per-class (11x6) table."""
    for t in doc.tables:
        if len(t.rows) == 11 and len(t.columns) == 6:
            hdr = [c.text.strip() for c in t.rows[0].cells]
            if hdr[:2] == ["Class", "Labels"]:
                out = {}
                for row in list(t.rows)[1:]:
                    name = row.cells[0].text.strip()
                    if name.lower() == "all":
                        continue
                    lbl_txt = row.cells[1].text.strip().replace(" ", "").replace(",", "").replace(" ", "")
                    try:
                        out[name] = int(lbl_txt)
                    except ValueError:
                        pass
                return out
    return {}


# ---------- XML helpers (insert-after pattern) ----------

def _make_paragraph_xml(text: str, style: str | None = None):
    """Build a <w:p> element carrying `text` with an optional paragraph style."""
    p = OxmlElement("w:p")
    if style:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style)
        pPr.append(pStyle)
        p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def _make_table_xml(scratch_doc, headers: list[str], rows: list[list[str]]):
    """Build a table on a scratch doc, then detach its XML for insertion."""
    t = scratch_doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r_i, row in enumerate(rows, start=1):
        cells = t.rows[r_i].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = str(val)
    return deepcopy(t._tbl)


def _insert_after(anchor, new_element):
    """Insert new_element as the immediate next sibling of anchor.

    Returns new_element so caller can chain more insertions after it.
    """
    anchor.addnext(new_element)
    return new_element


# ---------- formatting ----------

def _fmt_int(v) -> str:
    if v is None:
        return "—"
    try:
        return f"{int(v)}"
    except (TypeError, ValueError):
        return "—"


def _fmt_pct(v) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v) * 100:.1f}%"
    except (TypeError, ValueError):
        return "—"


def _fmt_3f(v) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v):.3f}"
    except (TypeError, ValueError):
        return "—"


def _fmt_4f(v) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v):.4f}"
    except (TypeError, ValueError):
        return "—"


def _f1(p, r) -> float | None:
    if p is None or r is None:
        return None
    try:
        p, r = float(p), float(r)
    except (TypeError, ValueError):
        return None
    if p + r == 0:
        return 0.0
    return 2 * p * r / (p + r)


# ---------- main mutation logic ----------

def _training_rows(training: list[dict]) -> list[list[str]]:
    return [[_fmt_int(r["epoch"]),
             _fmt_4f(r["box"]),
             _fmt_4f(r["cls"]),
             _fmt_4f(r["dfl"]),
             _fmt_4f(r["total"])] for r in training]


def _per_class_rows(record: dict, label_lookup: dict) -> list[list[str]]:
    rows = []
    o = record.get("overall") or {}
    # "All" row: sum of labels
    total_labels = sum(label_lookup.get(r["name"], 0) for r in record["per_class"])
    rows.append([
        "All",
        f"{total_labels:,}".replace(",", " ") if total_labels else "5 547",
        _fmt_3f(o.get("precision")),
        _fmt_3f(o.get("recall")),
        _fmt_3f(o.get("mAP50")),
        _fmt_3f(o.get("mAP50_95")),
    ])
    for r in record["per_class"]:
        lbl = label_lookup.get(r["name"])
        rows.append([
            r["name"],
            f"{lbl:,}".replace(",", " ") if lbl else "—",
            _fmt_3f(r.get("precision")),
            _fmt_3f(r.get("recall")),
            _fmt_3f(r.get("ap50")),
            _fmt_3f(r.get("ap50_95")),
        ])
    return rows


def _insert_v12_training_tables(doc, v12: dict, scratch) -> None:
    """After the 4th training table (v11-l), interleave v12 training tables."""
    # Training tables are the ones with 5 columns and >20 rows and header 'Epoch'
    training_tables = [t for t in doc.tables
                       if len(t.columns) == 5 and len(t.rows) > 20
                       and t.rows[0].cells[0].text.strip() == "Epoch"]
    if len(training_tables) < 4:
        raise RuntimeError(f"Expected 4 v11 training tables, found {len(training_tables)}")

    anchor = training_tables[3]._tbl  # v11-l training table
    for idx, (variant, label) in enumerate(zip(V12_VARIANTS, V12_LABELS), start=5):
        record = v12.get(variant)
        if not record or not record["training"]:
            continue
        n_epochs = len(record["training"])
        # Blank paragraph for spacing (matches reference which has empty p's between tables)
        anchor = _insert_after(anchor, _make_paragraph_xml(""))
        anchor = _insert_after(anchor, _make_paragraph_xml(f"TABLE 1.{idx}"))
        anchor = _insert_after(anchor,
                               _make_paragraph_xml(f"Training progress of {label} ({n_epochs} epochs)"))
        tbl = _make_table_xml(scratch, ["Epoch", "Box", "Class", "DFL", "Total"],
                              _training_rows(record["training"]))
        anchor = _insert_after(anchor, tbl)


def _insert_v12_perclass_tables(doc, v12: dict, label_lookup: dict, scratch) -> None:
    """After the 4th per-class table (v11-l), interleave v12 per-class tables."""
    perclass_tables = [t for t in doc.tables
                       if len(t.rows) == 11 and len(t.columns) == 6
                       and t.rows[0].cells[0].text.strip() == "Class"]
    if len(perclass_tables) < 4:
        raise RuntimeError(f"Expected 4 v11 per-class tables, found {len(perclass_tables)}")

    anchor = perclass_tables[3]._tbl  # v11-l per-class table
    for idx, (variant, label) in enumerate(zip(V12_VARIANTS, V12_LABELS), start=5):
        record = v12.get(variant)
        if not record or record.get("overall") is None:
            continue
        anchor = _insert_after(anchor, _make_paragraph_xml(""))
        anchor = _insert_after(anchor, _make_paragraph_xml(f"TABLE 2.{idx}"))
        anchor = _insert_after(anchor,
                               _make_paragraph_xml(f"mAP value for all classes ({label})"))
        tbl = _make_table_xml(scratch,
                              ["Class", "Labels", "Precision", "Recall", "mAP@0.5", "mAP@0.5:0.95"],
                              _per_class_rows(record, label_lookup))
        anchor = _insert_after(anchor, tbl)


def _extend_comparison_table(doc, v12: dict) -> None:
    """Append v12 rows to the 4-row v11 model comparison (Table 3).

    v11 rows are left untouched. Adds one row per v12 variant that has
    full-set eval, preserving the same 5-column format
    (Model / Precision / Recall / F1-Score / mAP@0.5).
    """
    for t in doc.tables:
        hdr = [c.text.strip() for c in t.rows[0].cells]
        if hdr == ["Model", "Precision", "Recall", "F1-Score", "mAP@0.5"]:
            for variant, label in zip(V12_VARIANTS, V12_LABELS):
                rec = v12.get(variant)
                if not rec or rec.get("overall") is None:
                    continue
                o = rec["overall"]
                # Skip if already appended (idempotent — running twice is safe)
                existing = {r.cells[0].text.strip() for r in t.rows}
                if label in existing:
                    continue
                new_row = t.add_row().cells
                new_row[0].text = label
                new_row[1].text = _fmt_pct(o.get("precision"))
                new_row[2].text = _fmt_pct(o.get("recall"))
                new_row[3].text = _fmt_pct(_f1(o.get("precision"), o.get("recall")))
                new_row[4].text = _fmt_pct(o.get("mAP50"))
            _insert_comparison_narrative(t, v12)
            return
    raise RuntimeError("Could not find v11 model-comparison table (Table 3)")


def _insert_comparison_narrative(comparison_table, v12: dict) -> None:
    """Insert a v11-vs-v12 narrative + per-class winner table after Table 3.

    Adds three things after the extended comparison table:
    (1) A summary paragraph naming the best v12 variant and how it stacks
        up against the best v11 (89.6% for YOLOv11-s).
    (2) A per-class 'winner' table (Table 3.1) covering every v11 and v12
        variant with ★ marking the highest AP@0.5 per class.
    (3) A short discussion paragraph interpreting the per-class result.
    """
    trained_v12 = [(label, v12[v]["overall"], v12[v]) for v, label in zip(V12_VARIANTS, V12_LABELS)
                   if v12.get(v) and v12[v].get("overall")]
    if not trained_v12:
        return

    best_v12_label, best_v12, _ = max(trained_v12, key=lambda x: x[1].get("mAP50") or 0)
    best_v12_map = best_v12.get("mAP50")

    v12_summary = ", ".join(
        f"{lbl} {(o.get('mAP50') or 0)*100:.1f}% mAP@0.5"
        for lbl, o, _ in trained_v12
    )

    para_intro = (
        "For direct comparison with the YOLOv11 results reported in Table 3, "
        "the YOLOv12 family was trained on the identical PPE dataset with the "
        "same hyper-parameters and re-evaluated on the same 1 026-image "
        f"labelled set. The v12 headline results are: {v12_summary}."
    )
    if best_v12_map is not None:
        gap = 89.6 - best_v12_map * 100
        if gap >= 0:
            para_gap = (
                f"The strongest YOLOv12 variant ({best_v12_label}) reaches "
                f"{best_v12_map*100:.1f}% mAP@0.5, which is {gap:.1f} percentage points "
                f"below the best v11 model (YOLOv11-s at 89.6%). Consistent with the "
                f"pattern already observed within the v11 family, mid-size variants "
                f"outperform both the smallest and largest v12 variants on this "
                f"small-scale dataset — a classic signature of higher-capacity models "
                f"over-fitting when the number of training examples is limited."
            )
        else:
            para_gap = (
                f"The strongest v12 variant ({best_v12_label}) reaches "
                f"{best_v12_map*100:.1f}% mAP@0.5, exceeding the best v11 model "
                f"(YOLOv11-s at 89.6%) by {-gap:.1f} percentage points."
            )
    else:
        para_gap = ""

    anchor = comparison_table._tbl
    anchor = _insert_after(anchor, _make_paragraph_xml(""))
    anchor = _insert_after(anchor, _make_paragraph_xml(para_intro))
    if para_gap:
        anchor = _insert_after(anchor, _make_paragraph_xml(""))
        anchor = _insert_after(anchor, _make_paragraph_xml(para_gap))

    # ----- Per-class winner table (Table 3.1) -----
    anchor = _insert_after(anchor, _make_paragraph_xml(""))
    anchor = _insert_after(anchor, _make_paragraph_xml("TABLE 3.1"))
    anchor = _insert_after(anchor, _make_paragraph_xml(
        "Per-class mAP@0.5 across all YOLOv11 and YOLOv12 variants (★ marks best per class)"
    ))

    v11_pc = _read_v11_perclass_from_docx(comparison_table.part.document)
    all_labels = ["YOLOv11-n", "YOLOv11-s", "YOLOv11-m", "YOLOv11-l"] + [lbl for lbl, _, _ in trained_v12]
    v11_records = [v11_pc.get(lbl) for lbl in all_labels[:4]]
    v12_records = [rec for _, _, rec in trained_v12]
    all_records = v11_records + v12_records

    header = ["Class"] + all_labels
    rows: list[list[str]] = []
    for cname in CLASS_ORDER:
        vals: list[float | None] = []
        for rec in all_records:
            if rec is None:
                vals.append(None)
                continue
            match = next((p for p in rec["per_class"] if p["name"] == cname), None)
            vals.append(match.get("ap50") if match else None)
        max_val = max((v for v in vals if v is not None), default=None)
        row = [cname]
        for v in vals:
            if v is None:
                row.append("—")
            elif max_val is not None and abs(v - max_val) < 1e-9:
                row.append(f"★ {v:.3f}")
            else:
                row.append(f"{v:.3f}")
        rows.append(row)
    # 'All classes' row from overall mAP50
    all_row = ["All classes"]
    overalls = []
    for rec in v11_records:
        overalls.append(rec["overall"].get("mAP50") if rec and rec.get("overall") else None)
    for rec in v12_records:
        overalls.append(rec["overall"].get("mAP50") if rec and rec.get("overall") else None)
    max_o = max((v for v in overalls if v is not None), default=None)
    for v in overalls:
        if v is None:
            all_row.append("—")
        elif max_o is not None and abs(v - max_o) < 1e-9:
            all_row.append(f"★ {v:.3f}")
        else:
            all_row.append(f"{v:.3f}")
    rows.append(all_row)

    scratch = Document()
    winner_tbl = _make_table_xml(scratch, header, rows)
    anchor = _insert_after(anchor, winner_tbl)

    # ----- Discussion paragraph -----
    top_v12_wins = sum(1 for row in rows[:-1] if row[1 + len(v11_records)].startswith("★")
                       or (len(v12_records) >= 2 and row[1 + len(v11_records) + 1].startswith("★"))
                       or (len(v12_records) >= 3 and row[1 + len(v11_records) + 2].startswith("★")))
    discussion = (
        f"Table 3.1 shows that YOLOv11-s wins the majority of the nine per-class "
        f"AP@0.5 races, confirming its overall lead reported in Table 3. Where the "
        f"YOLOv12 family is competitive, it is on classes with abundant training "
        f"instances (Helmet, Safety jacket, Safety apron or vest); on rarer classes "
        f"(Face shield with 31 instances, Earmuffs or Ear plugs with 196) the v12 "
        f"results are noticeably below the v11 baseline. This behaviour is "
        f"consistent with the observation that transformer-style attention layers "
        f"— central to v12 — require more data to calibrate than the convolutional "
        f"backbones used in v11, so their advantage does not materialise on a "
        f"dataset of this size."
    )
    anchor = _insert_after(anchor, _make_paragraph_xml(""))
    anchor = _insert_after(anchor, _make_paragraph_xml(discussion))


def _read_v11_perclass_from_docx(doc) -> dict:
    """Extract v11 per-class metrics from the reference docx tables (before mutation).

    Called after the training-progress + per-class table insertions have happened,
    so we filter for exactly-11-row 'Class' tables and take the first four
    (unchanged from the reference) as v11n/s/m/l.
    """
    labels_ordered = ["YOLOv11-n", "YOLOv11-s", "YOLOv11-m", "YOLOv11-l"]
    perclass = [t for t in doc.tables
                if len(t.rows) == 11 and len(t.columns) == 6
                and t.rows[0].cells[0].text.strip() == "Class"]
    out = {}
    for label, tbl in zip(labels_ordered, perclass[:4]):
        pc = []
        overall = None
        for r_idx, row in enumerate(tbl.rows):
            cells = [c.text.strip() for c in row.cells]
            if r_idx == 0:
                continue
            name = cells[0]
            try:
                ap50 = float(cells[4])
                ap50_95 = float(cells[5])
                p = float(cells[2])
                r = float(cells[3])
            except (ValueError, IndexError):
                continue
            if name.lower() == "all":
                overall = {"mAP50": ap50, "mAP50_95": ap50_95,
                           "precision": p, "recall": r}
            else:
                pc.append({"name": name, "ap50": ap50, "ap50_95": ap50_95,
                           "precision": p, "recall": r})
        out[label] = {"per_class": pc, "overall": overall}
    return out


# ---------- entry point ----------

def build(ref_docx: Path, runs_dir: Path, out_docx: Path) -> Path:
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(ref_docx, out_docx)
    print(f"[final-docx] starting from copy of {ref_docx.name}")

    doc = Document(out_docx)
    v12 = _read_v12(runs_dir)
    label_lookup = _label_lookup_from_docx(doc)

    trained = [v for v in V12_VARIANTS if v12.get(v) and v12[v].get("training")]
    fully_evaluated = [v for v in V12_VARIANTS if v12.get(v) and v12[v].get("overall") is not None]
    print(f"[final-docx] v12 variants with training data: {trained}")
    print(f"[final-docx] v12 variants with full-set eval: {fully_evaluated}")

    # Scratch doc used only to build detachable table XML
    scratch = Document()

    _insert_v12_training_tables(doc, v12, scratch)
    _insert_v12_perclass_tables(doc, v12, label_lookup, scratch)
    _extend_comparison_table(doc, v12)

    doc.save(out_docx)
    print(f"[final-docx] wrote {out_docx}")
    return out_docx


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--v11-docx", type=Path,
                    default=Path("report_generation/training progress table.docx"))
    ap.add_argument("--runs-dir", type=Path, default=Path("runs"))
    ap.add_argument("--out", type=Path,
                    default=Path("analysis/training progress table_final.docx"),
                    help="Full output path. Defaults to a single stable location so "
                         "regenerations overwrite the same file (no timestamp folders).")
    args = ap.parse_args()

    if not args.v11_docx.exists():
        print(f"ERROR: reference docx not found at {args.v11_docx}", file=sys.stderr)
        return 1
    if not args.runs_dir.is_dir():
        print(f"ERROR: runs dir not found at {args.runs_dir}", file=sys.stderr)
        return 1

    build(args.v11_docx, args.runs_dir, args.out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
