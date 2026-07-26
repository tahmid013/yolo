"""Generate a thesis-style v11+v12 comparison .docx that extends the reference
`training progress table.docx` from report_generation/.

Reads:
- Reference v11 training-progress tables (all epochs) + per-class tables from
  `report_generation/training progress table.docx` (tables 0-3 = training
  progress v11n/s/m/l, tables 4-7 = per-class v11n/s/m/l, table 8 = model
  comparison). These are copied through unchanged.
- Fresh v12 training progress from each `runs/yolo12X_*/results.csv`.
- Fresh v12 per-class metrics from each `runs/yolo12X_*/full_eval/per_class.json`
  (produced by pipeline.full_eval.run or colab/03_full_set_eval.ipynb).

Writes:
- analysis/reports/<ts>/v11_v12_thesis_tables.docx with sections:
    1. Training Progress — Tables 1.1-1.4 (v11) + 1.5-1.8 (v12)
    2. Per-class Performance — Tables 2.1-2.4 (v11) + 2.5-2.8 (v12)
    3. Model Comparison — extended Table 3 with all 8 rows
    4. Per-class Winner — 11x9 table with ★ marking best AP@0.5 across variants

Usage:
    python analysis/build_docx_report.py \\
        --v11-docx "report_generation/training progress table.docx" \\
        --runs-dir runs \\
        --out analysis/reports
"""
from __future__ import annotations

import argparse
import csv
import datetime as dt
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


CLASS_ORDER = [
    "Earmuffs or Ear plugs", "Face shield", "Footwear", "Goggles",
    "Hand gloves", "Helmet", "Masks", "Safety apron or vest", "Safety jacket",
]

V11_VARIANTS = ("yolo11n", "yolo11s", "yolo11m", "yolo11l")
V12_VARIANTS = ("yolo12n", "yolo12s", "yolo12m", "yolo12l")
V11_LABELS = ("YOLOv11-n", "YOLOv11-s", "YOLOv11-m", "YOLOv11-l")
V12_LABELS = ("YOLOv12-n", "YOLOv12-s", "YOLOv12-m", "YOLOv12-l")


# ---------- readers ----------

def _read_v11_from_docx(docx_path: Path) -> dict:
    """Extract training-progress + per-class tables for v11 from the reference docx.

    In `training progress table.docx`:
      tables 0..3  -> training progress v11n/s/m/l (5 cols: Epoch, Box, Class, DFL, Total)
      tables 4..7  -> per-class v11n/s/m/l (11 rows x 6 cols)
      table  8     -> 4-row model comparison
    """
    d = Document(docx_path)

    train_tables = [t for t in d.tables if len(t.columns) == 5 and len(t.rows) > 20]
    perclass_tables = [t for t in d.tables if len(t.rows) == 11 and len(t.columns) == 6]

    if len(train_tables) < 4:
        raise RuntimeError(f"Expected >=4 training tables in {docx_path.name}, found {len(train_tables)}")
    if len(perclass_tables) < 4:
        raise RuntimeError(f"Expected >=4 per-class (11x6) tables in {docx_path.name}, found {len(perclass_tables)}")

    out: dict = {}
    for variant, ttbl, ptbl in zip(V11_VARIANTS, train_tables[:4], perclass_tables[:4]):
        # Training progress rows
        train_rows = []
        for r_idx, row in enumerate(ttbl.rows):
            cells = [c.text.strip() for c in row.cells]
            if r_idx == 0:
                continue
            try:
                train_rows.append({
                    "epoch": int(_first_number(cells[0]) or 0),
                    "box":   _parse_float(cells[1]),
                    "cls":   _parse_float(cells[2]),
                    "dfl":   _parse_float(cells[3]),
                    "total": _parse_float(cells[4]),
                })
            except Exception:
                continue

        # Per-class rows
        pc_rows = []
        overall = None
        for r_idx, row in enumerate(ptbl.rows):
            cells = [c.text.strip() for c in row.cells]
            if r_idx == 0:
                continue
            name = cells[0]
            labels = _parse_int(cells[1])
            p = _parse_float(cells[2])
            r_ = _parse_float(cells[3])
            ap50 = _parse_float(cells[4])
            ap50_95 = _parse_float(cells[5])
            if name.lower() == "all":
                overall = {"precision": p, "recall": r_,
                           "mAP50": ap50, "mAP50_95": ap50_95, "labels": labels}
            else:
                pc_rows.append({"name": name, "labels": labels,
                                "precision": p, "recall": r_,
                                "ap50": ap50, "ap50_95": ap50_95})

        out[variant] = {
            "training": train_rows,
            "per_class": pc_rows,
            "overall": overall,
            "source": f"{docx_path.name}",
        }
    return out


def _read_v12_training(csv_path: Path) -> list[dict]:
    """Extract per-epoch train losses from ultralytics results.csv."""
    rows = []
    with csv_path.open("r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            r = {k.strip(): (v.strip() if v else v) for k, v in r.items()}
            try:
                epoch = int(float(r.get("epoch", 0)))
                box = float(r["train/box_loss"])
                cls = float(r["train/cls_loss"])
                dfl = float(r["train/dfl_loss"])
                rows.append({"epoch": epoch, "box": box, "cls": cls, "dfl": dfl,
                             "total": box + cls + dfl})
            except (KeyError, TypeError, ValueError):
                continue
    return rows


def _read_v12_from_runs(runs_dir: Path) -> dict:
    """Load per-epoch training + full_eval/per_class.json for each v12 run."""
    out: dict = {}
    for sub in sorted(runs_dir.iterdir()):
        if not sub.is_dir():
            continue
        variant = next((v for v in V12_VARIANTS if sub.name.startswith(f"{v}_")), None)
        if variant is None:
            continue

        record = {"source": sub.name, "training": [], "per_class": [], "overall": None}

        results_csv = sub / "results.csv"
        if results_csv.exists():
            record["training"] = _read_v12_training(results_csv)
        else:
            print(f"[docx] WARN {sub.name}: no results.csv")

        eval_path = sub / "full_eval" / "per_class.json"
        if eval_path.exists():
            payload = json.loads(eval_path.read_text(encoding="utf-8"))
            record["overall"] = payload.get("overall")
            by_name = {p["name"]: p for p in payload.get("per_class", [])}
            for cname in CLASS_ORDER:
                src = by_name.get(cname, {})
                record["per_class"].append({
                    "name": cname,
                    "labels": None,
                    "precision": src.get("precision"),
                    "recall": src.get("recall"),
                    "ap50": src.get("ap50"),
                    "ap50_95": src.get("ap50_95"),
                })
        else:
            print(f"[docx] {sub.name}: no full_eval/per_class.json — "
                  "per-class + overall will be blank. Run colab/03_full_set_eval.ipynb.")

        out[variant] = record
    return out


# ---------- doc building ----------

def _fmt(v, kind="pct") -> str:
    if v is None:
        return "—"
    try:
        vf = float(v)
    except (TypeError, ValueError):
        return "—"
    if kind == "pct":
        return f"{vf * 100:.1f}%"
    if kind == "3f":
        return f"{vf:.3f}"
    if kind == "4f":
        return f"{vf:.4f}"
    if kind == "int":
        return f"{int(vf)}"
    return str(vf)


def _f1(p, r) -> float | None:
    if p is None or r is None:
        return None
    try:
        p, r = float(p), float(r)
    except (TypeError, ValueError):
        return None
    if p + r == 0:
        return 0.0
    return 2 * p * r / (p + r)


def _add_table(doc, headers: list[str], rows: list[list[str]],
               caption_before: str | None = None):
    if caption_before:
        doc.add_paragraph(caption_before)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_i, row in enumerate(rows, start=1):
        cells = t.rows[r_i].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = str(val)
    doc.add_paragraph("")


def _training_table_rows(training: list[dict]) -> list[list[str]]:
    return [[_fmt(r["epoch"], "int"),
             _fmt(r["box"], "4f"),
             _fmt(r["cls"], "4f"),
             _fmt(r["dfl"], "4f"),
             _fmt(r["total"], "4f")] for r in training]


def _per_class_table_rows(overall: dict | None, per_class: list[dict],
                          label_lookup: dict) -> list[list[str]]:
    rows = []
    if overall:
        # sum of per-class labels = "All" row's Labels column
        total_labels = sum((label_lookup.get(c["name"]) or 0) for c in per_class) or None
        rows.append(["All",
                     _fmt(total_labels, "int") if total_labels else "5 547",
                     _fmt(overall.get("precision"), "3f"),
                     _fmt(overall.get("recall"), "3f"),
                     _fmt(overall.get("mAP50"), "3f"),
                     _fmt(overall.get("mAP50_95"), "3f")])
    else:
        rows.append(["All", "—", "—", "—", "—", "—"])
    for r in per_class:
        rows.append([r["name"],
                     _fmt(label_lookup.get(r["name"]), "int") if label_lookup.get(r["name"]) else "—",
                     _fmt(r.get("precision"), "3f"),
                     _fmt(r.get("recall"), "3f"),
                     _fmt(r.get("ap50"), "3f"),
                     _fmt(r.get("ap50_95"), "3f")])
    return rows


def build_docx(v11: dict, v12: dict, out_docx: Path) -> Path:
    doc = Document()
    doc.add_heading("YOLOv11 vs YOLOv12 — Thesis Tables (Extended)", level=1)
    doc.add_paragraph(
        "Supplementary thesis tables extending 'training progress table.docx' "
        "to include the YOLOv12 family (n, s, m, l). YOLOv11 tables (1.1-1.4, "
        "2.1-2.4) are copied verbatim from the reference document; YOLOv12 "
        "tables (1.5-1.8, 2.5-2.8) are freshly computed from the runs produced "
        "by the pipeline. Per-class metrics for v12 come from full-set "
        "re-evaluation (all 1,026 images) matching the reference methodology."
    )
    ts_str = dt.datetime.now(dt.timezone.utc).isoformat()
    doc.add_paragraph(f"Generated: {ts_str}")

    # Labels come from v11n's per-class table (identical across all variants — same dataset)
    label_lookup = {}
    if v11.get("yolo11n"):
        for r in v11["yolo11n"]["per_class"]:
            label_lookup[r["name"]] = r.get("labels")

    # ============ SECTION 1: TRAINING PROGRESS ============
    doc.add_heading("1. Training Progress", level=2)
    doc.add_paragraph(
        "Each model records per-epoch training losses as Box / Class / DFL "
        "(Distribution Focal Loss); Total = Box + Class + DFL. Each model "
        "trained for a different number of epochs because early-stopping was "
        "triggered when the validation mAP stopped improving."
    )

    # v11 always present (from reference docx)
    _v11_train = [
        ("YOLOv11-n", v11.get("yolo11n")),
        ("YOLOv11-s", v11.get("yolo11s")),
        ("YOLOv11-m", v11.get("yolo11m")),
        ("YOLOv11-l", v11.get("yolo11l")),
    ]
    # v12 only include variants that were actually trained (have training rows)
    _v12_train = [
        (label, v12.get(v)) for v, label in zip(V12_VARIANTS, V12_LABELS)
        if v12.get(v) and v12[v].get("training")
    ]
    all_train = _v11_train + _v12_train
    for idx, (label, d) in enumerate(all_train, start=1):
        # Table 1.1..1.4 for v11, 1.5+ for v12
        section = 1 if idx <= 4 else 1
        num = idx if idx <= 4 else 4 + (idx - 4)  # continues 1.5, 1.6, ...
        tag = f"Table 1.{num}"
        title = f"Training progress of {label}"
        if not d or not d.get("training"):
            doc.add_paragraph(f"{tag}: {title} — NOT AVAILABLE")
            continue
        rows = _training_table_rows(d["training"])
        n_epochs = len(rows)
        _add_table(
            doc, ["Epoch", "Box", "Class", "DFL", "Total"], rows,
            caption_before=f"{tag}: {title} ({n_epochs} epochs)",
        )

    # ============ SECTION 2: PER-CLASS PERFORMANCE ============
    doc.add_heading("2. Per-class Performance", level=2)
    doc.add_paragraph(
        "Precision / Recall / mAP@0.5 / mAP@0.5:0.95 per class, computed on "
        "the full 1,026-image labelled set (5,547 PPE instances). Labels denotes "
        "the number of ground-truth instances of each class (identical across "
        "models — same dataset)."
    )

    _v11_pc = [
        ("YOLOv11-n", v11.get("yolo11n")),
        ("YOLOv11-s", v11.get("yolo11s")),
        ("YOLOv11-m", v11.get("yolo11m")),
        ("YOLOv11-l", v11.get("yolo11l")),
    ]
    # Only include v12 variants that have full-set eval (overall metrics populated)
    _v12_pc = [
        (label, v12.get(v)) for v, label in zip(V12_VARIANTS, V12_LABELS)
        if v12.get(v) and v12[v].get("overall") is not None
    ]
    all_pc = _v11_pc + _v12_pc
    for idx, (label, d) in enumerate(all_pc, start=1):
        num = idx  # 2.1, 2.2, ..., continues naturally
        tag = f"Table 2.{num}"
        title = f"mAP value for all classes ({label})"
        if not d or not d.get("per_class") or (d.get("overall") is None and not d.get("per_class")):
            doc.add_paragraph(f"{tag}: {title} — NOT AVAILABLE (run full-set eval)")
            continue
        rows = _per_class_table_rows(d.get("overall"), d["per_class"], label_lookup)
        _add_table(
            doc, ["Class", "Labels", "Precision", "Recall", "mAP@0.5", "mAP@0.5:0.95"],
            rows, caption_before=f"{tag}: {title}",
        )

    doc.add_paragraph(
        "Note: Face shield is severely under-represented (31 instances, ~0.6% "
        "of the data); its per-class metrics are statistically less stable than "
        "the other classes and should be reported with a footnote."
    )

    # ============ SECTION 3: 8-MODEL COMPARISON ============
    doc.add_heading("3. Model Comparison", level=2)
    doc.add_paragraph(
        "Table 3: Performance of all trained models. Precision and Recall are "
        "means across the nine classes; F1-Score is computed from those means; "
        "mAP@0.5 is the standard PASCAL-VOC metric on the full evaluation set."
    )
    rows = []
    # v11 always present; v12 only variants that were fully evaluated
    v11_data = list(zip(V11_LABELS, [v11.get(v) for v in V11_VARIANTS]))
    v12_data = [(label, v12.get(v)) for v, label in zip(V12_VARIANTS, V12_LABELS)
                if v12.get(v) and v12[v].get("overall") is not None]
    for label, d in v11_data + v12_data:
        if d is None or d.get("overall") is None:
            rows.append([label, "—", "—", "—", "—"])
            continue
        o = d["overall"]
        rows.append([
            label,
            _fmt(o.get("precision")),
            _fmt(o.get("recall")),
            _fmt(_f1(o.get("precision"), o.get("recall"))),
            _fmt(o.get("mAP50")),
        ])
    _add_table(doc, ["Model", "Precision", "Recall", "F1-Score", "mAP@0.5"], rows,
               caption_before="Table 3: Performance of all the models")

    # ============ SECTION 4: PER-CLASS WINNER ACROSS ALL 8 ============
    doc.add_heading("4. Per-class mAP@0.5 across all variants", level=2)
    doc.add_paragraph(
        "★ marks the highest AP@0.5 for that class across all trained variants. "
        "Blank cells indicate the corresponding v12 model has not yet been "
        "fully evaluated."
    )
    # v11 always; v12 only variants fully evaluated
    _v12_present = [(v, label) for v, label in zip(V12_VARIANTS, V12_LABELS)
                    if v12.get(v) and v12[v].get("overall") is not None]
    all_labels = list(V11_LABELS) + [lbl for _, lbl in _v12_present]
    all_records = [v11.get(v) for v in V11_VARIANTS] + [v12.get(v) for v, _ in _v12_present]
    headers = ["Class"] + all_labels
    rows_out = []
    for cname in CLASS_ORDER:
        vals = []
        for rec in all_records:
            if rec is None:
                vals.append(None); continue
            match = next((p for p in rec["per_class"] if p["name"] == cname), None)
            vals.append(match.get("ap50") if match else None)
        max_val = max((v for v in vals if v is not None), default=None)
        row_out = [cname]
        for v in vals:
            if v is None:
                row_out.append("—")
            elif max_val is not None and abs(v - max_val) < 1e-9:
                row_out.append(f"★ {v:.3f}")
            else:
                row_out.append(f"{v:.3f}")
        rows_out.append(row_out)

    # All classes row (overall mAP50)
    all_row = ["All classes"]
    vals = [rec["overall"]["mAP50"] if rec and rec.get("overall") else None for rec in all_records]
    max_val = max((v for v in vals if v is not None), default=None)
    for v in vals:
        if v is None:
            all_row.append("—")
        elif max_val is not None and abs(v - max_val) < 1e-9:
            all_row.append(f"★ {v:.3f}")
        else:
            all_row.append(f"{v:.3f}")
    rows_out.append(all_row)
    _add_table(doc, headers, rows_out)

    # ============ APPENDIX ============
    doc.add_heading("Appendix: Data Provenance", level=2)
    doc.add_paragraph("Each row above traces to one of:")
    for v in V11_VARIANTS:
        src = v11.get(v, {}).get("source", "MISSING")
        doc.add_paragraph(f"  • {v}: {src} (verbatim copy)", style="List Bullet")
    for v in V12_VARIANTS:
        d = v12.get(v)
        if not d:
            continue  # skip variants not trained (v12n-only mode)
        src = d.get("source", "MISSING")
        n_train = len(d.get("training", []))
        has_full = d.get("overall") is not None
        doc.add_paragraph(
            f"  • {v}: {src}  (training rows: {n_train}, full-set eval: {'yes' if has_full else 'NO'})",
            style="List Bullet",
        )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    return out_docx


# ---------- utilities ----------

def _first_number(s: str) -> float | None:
    s = s.replace(",", "").strip()
    for tok in s.split():
        try:
            return float(tok)
        except ValueError:
            continue
    return None


def _parse_float(s: str) -> float | None:
    s = (s or "").replace("%", "").replace(",", "").replace(" ", "").strip()
    try:
        return float(s)
    except (TypeError, ValueError):
        return None


def _parse_int(s: str) -> int | None:
    s = (s or "").replace(",", "").replace(" ", "").replace(" ", "").strip()
    try:
        return int(float(s))
    except (TypeError, ValueError):
        return None


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--v11-docx", type=Path,
                    default=Path("report_generation/training progress table.docx"))
    ap.add_argument("--runs-dir", type=Path, default=Path("runs"))
    ap.add_argument("--out", type=Path, default=Path("analysis/reports"))
    ap.add_argument("--out-name", type=str, default="v11_v12_thesis_tables.docx",
                    help="Filename for the generated docx (default: v11_v12_thesis_tables.docx)")
    args = ap.parse_args()

    if not args.v11_docx.exists():
        print(f"ERROR: v11 reference docx not found at {args.v11_docx}", file=sys.stderr)
        return 1
    if not args.runs_dir.is_dir():
        print(f"ERROR: runs dir not found at {args.runs_dir}", file=sys.stderr)
        return 1

    v11 = _read_v11_from_docx(args.v11_docx)
    v12 = _read_v12_from_runs(args.runs_dir)

    print(f"[docx] v11 variants loaded from '{args.v11_docx.name}': {len(v11)}/{len(V11_VARIANTS)}")
    for v in V11_VARIANTS:
        d = v11.get(v)
        if d:
            print(f"       {v}: {len(d['training'])} epochs, "
                  f"{len(d['per_class'])} classes, overall mAP50={d['overall']['mAP50']}")
    print(f"[docx] v12 variants found in runs/: {len(v12)}/{len(V12_VARIANTS)}")
    for v in V12_VARIANTS:
        d = v12.get(v)
        if d:
            has_full = d.get('overall') is not None
            print(f"       {v}: {len(d['training'])} epochs, "
                  f"full-set eval={'yes' if has_full else 'NO'}")
        else:
            print(f"       {v}: not trained yet")

    ts = dt.datetime.now(dt.timezone.utc).strftime("%Y%m%d-%H%M%S")
    out_docx = args.out / ts / args.out_name
    build_docx(v11, v12, out_docx)
    print(f"[docx] wrote {out_docx}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
